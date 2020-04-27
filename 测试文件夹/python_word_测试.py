#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2020/3/29 21:42
# @Author  : TanLHHH
# @Site    : 
# @File    : python_word_测试.py
# @Software: PyCharm

from docx import Document
document = Document()
style = document.styles['Normal']
document.add_heading('携程网以“湖北”为关键词的“目的地参团”旅游线路',level=0)


n1 = document.add_paragraph('目的地参团--跟团游旅游线路：',style='List Number')
n3 = document.add_paragraph('目的地参团--半自助游旅游线路：',style='List Number')
n4 = document.add_paragraph('目的地参团--私家团游旅游线路：',style='List Number')

text = """
({0}) {1}
    价格：{2}元/人
    行程：{3}
""".format()
n1.add_run(text=text)

document.save('demo.docx')