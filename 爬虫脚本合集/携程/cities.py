#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2020/3/30 02:06
# @Author  : TanLHHH
# @Site    : 
# @File    : 重庆.py
# @Software: PyCharm

city= ['重庆','西安','南京','杭州','厦门','成都','深圳','广州','昆明','长沙','沈阳','武汉']

import requests
import random
import re
import time
import json
# user_agent列表
user_agent_list=[
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.119 Safari/537.36",
    "Mozilla/5.0 (Windows NT 5.1; U; en; rv:1.8.1) Gecko/20061208 Firefox/2.0.0 Opera 9.50",
    "Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; en) Opera 9.50",
    "Mozilla/5.0 (Windows NT 6.1; WOW64; rv:34.0) Gecko/20100101 Firefox/34.0",
    "Mozilla/5.0 (X11; U; Linux x86_64; zh-CN; rv:1.9.2.10) Gecko/20100922 Ubuntu/10.10 (maverick) Firefox/3.6.10",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/534.57.2 (KHTML, like Gecko) Version/5.1.7 Safari/534.57.2",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.71 Safari/537.36",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.11 (KHTML, like Gecko) Chrome/23.0.1271.64 Safari/537.11",
    "Mozilla/5.0 (Windows; U; Windows NT 6.1; en-US) AppleWebKit/534.16 (KHTML, like Gecko) Chrome/10.0.648.133 Safari/534.16",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/30.0.1599.101 Safari/537.36",
    "Mozilla/5.0 (Windows NT 6.1; WOW64; Trident/7.0; rv:11.0) like Gecko",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.11 TaoBrowser/2.0 Safari/536.11",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/21.0.1180.71 Safari/537.1 LBBROWSER",
    "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; WOW64; Trident/5.0; SLCC2; .NET CLR 2.0.50727; .NET CLR 3.5.30729; .NET CLR 3.0.30729; Media Center PC 6.0; .NET4.0C; .NET4.0E; LBBROWSER) ",
    "Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; QQDownload 732; .NET4.0C; .NET4.0E; LBBROWSER)",
    "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; WOW64; Trident/5.0; SLCC2; .NET CLR 2.0.50727; .NET CLR 3.5.30729; .NET CLR 3.0.30729; Media Center PC 6.0; .NET4.0C; .NET4.0E; QQBrowser/7.0.3698.400)",
    "Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; QQDownload 732; .NET4.0C; .NET4.0E)",
    "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.84 Safari/535.11 SE 2.X MetaSr 1.0",
    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Trident/4.0; SV1; QQDownload 732; .NET4.0C; .NET4.0E; SE 2.X MetaSr 1.0)",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Maxthon/4.4.3.4000 Chrome/30.0.1599.101 Safari/537.36",
    "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/38.0.2125.122 UBrowser/4.0.3214.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:65.0) Gecko/20100101 Firefox/65.0",
    "Mozilla/5.0 (Macintosh; U; Intel Mac OS X 10_6_8; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50",
    "Mozilla/5.0 (Windows; U; Windows NT 6.1; en-us) AppleWebKit/534.50 (KHTML, like Gecko) Version/5.1 Safari/534.50",
    "Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; Trident/5.0",
    "Mozilla/4.0 (compatible; MSIE 8.0; Windows NT 6.0; Trident/4.0)",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10.6; rv:2.0.1) Gecko/20100101 Firefox/4.0.1",
    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; The World)",
    "Mozilla/5.0 (Windows NT 6.1; rv:2.0.1) Gecko/20100101 Firefox/4.0.1",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_0) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11",
    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Maxthon 2.0)",
    "Opera/9.80 (Android 2.3.4; Linux; Opera Mobi/build-1107180945; U; en-GB) Presto/2.8.149 Version/11.10",
    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; 360SE)",
    "MQQBrowser/26 Mozilla/5.0 (Linux; U; Android 2.3.7; zh-cn; MB200 Build/GRJ22; CyanogenMod-7) AppleWebKit/533.1 (KHTML, like Gecko) Version/4.0 Mobile Safari/533.1",
    "Mozilla/5.0 (Linux; U; Android 2.3.7; en-us; Nexus One Build/FRF91) AppleWebKit/533.1 (KHTML, like Gecko) Version/4.0 Mobile Safari/533.1",
    "Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Trident/4.0; SE 2.X MetaSr 1.0; SE 2.X MetaSr 1.0; .NET CLR 2.0.50727; SE 2.X MetaSr 1.0)",
]

from docx import Document


def document_init(city):
    global document
    document = Document()
    global n1,n3,n4
    document.add_heading('携程网以“湖北”为目的地的“出发地参团”旅游线路', level=0)
    n1 = document.add_paragraph('跟团游旅游线路：', style='List Number')
    n3 = document.add_paragraph('半自助游旅游线路：', style='List Number')
    n4 = document.add_paragraph('私家团游旅游线路：', style='List Number')
    document.save('{}出发.docx'.format(city))

    global n1_num , n3_num , n4_num
    n1_num = 1
    n3_num = 1
    n4_num = 1


def get_proxy():
    return requests.get("http://192.144.162.135:5010/get/").json()

#得到列表页html
def get_detail_url_list(city):
    proxy = get_proxy().get("proxy")
    headers={
            'User-Agent':random.choice(user_agent_list),
            'Connection': 'keep-alive',
            'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9'
    }
    #try:
    url = "https://vacations.ctrip.com/list/grouptravel/d-hubei-100067.html?cityname={0}&salecity=1&startcity=1&filter={1}&p={2}"

    # n1：跟团游 n3：半自助 n4：私家团
    filter_list = ['n1','n3','n4']
    for i in range(0,len(filter_list)):
        print("------------------当前正在爬取",filter_list[i],"----------------------------")
        p = 1
        while 1:
            res = requests.get(url=url.format(city,filter_list[i],str(p)),headers=headers,proxies={"http": "http://{}".format(proxy)})

            print("cur_url:",url.format(city,filter_list[i],str(p)))
            pat = '"products":(.*?),"crumbs":'
            products_info = re.compile(pat, re.S).findall(res.content.decode('utf-8'))
            if len(products_info) == 0:
                break

            #  详情页
            detail_url_pat = '"detailUrl":"(.*?)",'
            detail_url_list = re.compile(detail_url_pat, re.S).findall(products_info[0])

            #  产品ID
            products_id_pat = '"id":(.*?),'
            products_id_list_before = re.compile(products_id_pat, re.S).findall(products_info[0])

            products_id_list = []
            for id in products_id_list_before:
                if id not in products_id_list:
                    products_id_list.append(id)

            if len(detail_url_list) == 0:
                break
            p = p + 1

            parsing_detail_url_list(detail_url_list,products_id_list,filter_list[i],city)
            time.sleep(2)

    # except Exception as e:
    #     print(e)

def parsing_detail_url_list(detail_url_list,products_id_list,filter,city):
    url = 'https://vacations.ctrip.com/tour/restapi/online/12447/ProductDetailTimingV5?_fxpcqlniredt=09031174410487229982'
    data = {"ChannelCode": 0,
            "PlatformId": 4,
            "Version": "80400",
            "Locale": "zh-CN",
            "head": {"cid": "09031174410487229982", "ctok": "", "cver": "1.0", "lang": "01", "sid": "8888",
                     "syscode": "09", "auth": "", "extension": []},
            "DepartureCityId": 477,
            "QueryNode": {
                "IsTravelIntroductionInfo":
                "true", "IsBookInfo": "true",
                "IsBasicExtendInfo": "true",
                "IsSegmentInfo": "true"},
            "contentType": "json"}
    Header = {
        'Content-Type': 'application/json',
        'User-Agent':random.choice(user_agent_list),
    }
    #try:
    for i in range(0,len(products_id_list)):
        print("len(products_id_list)：",len(products_id_list),"当前：",i)
        data["ProductId"] = products_id_list[i]
        proxy = get_proxy().get("proxy")
        header = {
            'User-Agent':random.choice(user_agent_list)
        }
        cur_url = "https:" + str(detail_url_list[i])

        print("cur_url:",cur_url)
        res = requests.get(cur_url,headers=header,proxies={"http": "http://{}".format(proxy)})
        title_pat = '<title>(.*?)·.*?</title>'
        title = re.compile(title_pat,re.S).findall(res.content.decode('utf-8'))
        print("当前title为",title[0])

        price_pat = '<dfn>¥</dfn><em>(.*?)</em>'
        price = re.compile(price_pat,re.S).findall(res.content.decode('utf-8'))
        res = requests.post(url, data=json.dumps(data), headers=Header,proxies={"http": "http://{}".format(proxy)})
        pat = '"GSScenicSpotID":.*?,"Name":"(.*?)","PreName":'
        point_list = re.compile(pat,re.S).findall(res.content.decode('utf-8'))
        print(point_list)

        if len(price) == 0:
            print("price is None")
            price = 868
            write_to_world(filter, title[0], price, point_list,city)
        else:
            print("当前价格为", price[0], "元/人")
            write_to_world(filter,title[0],price[0],point_list,city)
        time.sleep(2)
    #except Exception as e:
        #print(e)

def write_to_world(filter,title,price,point_list,city):
    route = """"""
    print("当前filter为：",filter)

    global n1_num,n3_num,n4_num
    global n1,n3,n4

    for i in range(0,len(point_list)-1):
        route = route + str(point_list[i]) +'---'
    route += point_list[-1]
    if filter == 'n1':
        text = """
({0}) {1}
    价格：{2}元/人
    行程：{3}
""".format(n1_num,title,price,route)
        n1.add_run(text=text)
        n1_num += 1
        print("n1_add")
    elif filter == 'n3':
        text = """
({0}) {1}
    价格：{2}元/人
    行程：{3}
""".format(n3_num,title,price,route)
        n3.add_run(text=text)
        n3_num += 1
        print("n3_add")
    else:
        text = """
({0}) {1}
    价格：{2}元/人
    行程：{3}
""".format(n4_num,title,price,route)
        n4.add_run(text=text)
        n4_num += 1
        print("n4_add")
    document.save('{}出发.docx'.format(city))
    print("文件已经成功写入word文档")

def main():
    try:
        for i in range(0,len(city)):
            document_init(city[i])
            get_detail_url_list(city[i])
    except Exception as e:
        print(e)

if __name__ == '__main__':
    main()


